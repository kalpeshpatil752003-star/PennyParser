import os
import io
import logging

logger = logging.getLogger("extraction")

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None


def extract_pages(file_path: str) -> list[dict]:
    """Returns [{ "page_number": int, "text": str }, ...] for PDF documents."""
    if fitz is None:
        return extract_txt(file_path)
    doc = fitz.open(file_path)
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text().strip()
        
        # Scanned page detection: text layer is empty or negligible, but page has image content
        if len(text) < 10:
            images = page.get_images()
            if images:
                ocr_success = False
                try:
                    import pytesseract
                    from PIL import Image
                    pix = page.get_pixmap(dpi=150)
                    img = Image.open(io.BytesIO(pix.tobytes()))
                    ocr_text = pytesseract.image_to_string(img).strip()
                    if ocr_text:
                        text = f"[OCR Extracted Page {i + 1}]\n" + ocr_text
                        ocr_success = True
                except Exception as e:
                    logger.debug(f"OCR not available for scanned page {i + 1}: {e}")

                if not ocr_success:
                    text = f"[Notice: Page {i + 1} contains scanned/image content with no embedded text layer.]"

        if text:
            pages.append({"page_number": i + 1, "text": text})
    doc.close()
    return pages


def extract_docx(file_path: str) -> list[dict]:
    """Extracts text and structured tables from DOCX files, grouping into logical pages."""
    try:
        import docx
        doc = docx.Document(file_path)
        content_parts = []

        # Extract paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                content_parts.append(p.text.strip())

        # Extract tables in text format
        for t_idx, table in enumerate(doc.tables):
            table_lines = [f"\n[Table {t_idx + 1}]"]
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    table_lines.append(" | ".join(cells))
            if len(table_lines) > 1:
                content_parts.append("\n".join(table_lines))

        combined = "\n\n".join(content_parts)
    except ImportError:
        combined = ""
    except Exception as e:
        logger.warning(f"Error extracting DOCX: {e}")
        combined = ""

    if not combined:
        return extract_txt(file_path)

    return _split_into_logical_pages(combined)


def extract_txt(file_path: str) -> list[dict]:
    """Extracts text from TXT files, grouping lines into logical pages of ~2000 chars."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read().strip()
    if not text:
        return []
    return _split_into_logical_pages(text)


def _split_into_logical_pages(text: str, page_size: int = 2000) -> list[dict]:
    pages = []
    start = 0
    page_num = 1
    while start < len(text):
        end = start + page_size
        chunk = text[start:end].strip()
        if chunk:
            pages.append({"page_number": page_num, "text": chunk})
            page_num += 1
        start = end
    return pages


def extract_document(file_path: str, file_type: str) -> list[dict]:
    """
    Unified extraction dispatcher.
    Supports PDF, DOCX, TXT.
    Returns normalized list: [{ "page_number": int, "text": str }, ...]
    """
    file_type_upper = file_type.upper() if file_type else ""
    ext = os.path.splitext(file_path)[1].lower()

    if file_type_upper == "PDF" or ext == ".pdf":
        return extract_pages(file_path)
    elif file_type_upper == "DOCX" or ext == ".docx":
        return extract_docx(file_path)
    elif file_type_upper == "TXT" or ext == ".txt":
        return extract_txt(file_path)
    else:
        try:
            return extract_pages(file_path)
        except Exception:
            return extract_txt(file_path)